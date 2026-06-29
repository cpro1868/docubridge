from __future__ import annotations

from collections.abc import Iterable, Mapping
from pathlib import Path

from docx.document import Document as DocumentType
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt

from docubridge.adapters.docx_adapter import create_document
from docubridge.core.asset_resolver import resolve_image_path
from docubridge.core.layout_intent import NumberingIntent, ParagraphLayoutIntent
from docubridge.core.nodes import CodeBlockNode, HeadingNode, HorizontalRuleNode, ImageBlockNode, ListNode, ParagraphNode, QuoteNode, TableNode


def _render_inline_text(inlines: Iterable[object]) -> str:
    return "".join(getattr(inline, "text", "") for inline in inlines)


def _apply_run_style(run, properties: Mapping[str, object]) -> None:
    font_ascii = properties.get("font_ascii") or properties.get("font_name")
    if isinstance(font_ascii, str) and font_ascii:
        run.font.name = font_ascii
    font_size = properties.get("font_size")
    if isinstance(font_size, (int, float)):
        run.font.size = Pt(float(font_size))
    font_color = properties.get("font_color")
    if isinstance(font_color, str) and font_color:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor.from_string(font_color)
    if properties.get("bold") is True:
        run.bold = True
    if properties.get("italic") is True:
        run.italic = True
    if properties.get("strike") is True:
        run.font.strike = True
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for key, attr_name in (
        ("font_ascii", "ascii"),
        ("font_hansi", "hAnsi"),
        ("font_east_asia", "eastAsia"),
        ("font_cs", "cs"),
    ):
        value = properties.get(key)
        if isinstance(value, str) and value:
            r_fonts.set(qn(f"w:{attr_name}"), value)


def _add_inline_runs(paragraph, inlines: Iterable[object], *, properties: Mapping[str, object] | None = None) -> None:
    effective_properties = properties or {}
    for inline in inlines:
        href = getattr(inline, "href", None)
        if href:
            part = paragraph.part
            r_id = part.relate_to(href, RT.HYPERLINK, is_external=True)
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), r_id)
            run = OxmlElement("w:r")
            r_pr = OxmlElement("w:rPr")
            if getattr(inline, "bold", False):
                r_pr.append(OxmlElement("w:b"))
            if getattr(inline, "italic", False):
                r_pr.append(OxmlElement("w:i"))
            if getattr(inline, "strike", False):
                r_pr.append(OxmlElement("w:strike"))
            if getattr(inline, "code", False):
                r_fonts = OxmlElement("w:rFonts")
                r_fonts.set(qn("w:ascii"), "Consolas")
                r_fonts.set(qn("w:hAnsi"), "Consolas")
                r_pr.append(r_fonts)
            if len(r_pr):
                run.append(r_pr)
            text = OxmlElement("w:t")
            text.text = getattr(inline, "text", "")
            run.append(text)
            hyperlink.append(run)
            paragraph._p.append(hyperlink)
            continue
        run = paragraph.add_run(getattr(inline, "text", ""))
        if getattr(inline, "bold", False):
            run.bold = True
        if getattr(inline, "italic", False):
            run.italic = True
        if getattr(inline, "strike", False):
            run.font.strike = True
        if getattr(inline, "code", False):
            run.font.name = "Consolas"
        else:
            _apply_run_style(run, effective_properties)


def render_missing_image_placeholder(path: str) -> str:
    return f"[Image not found: {path}]"


def _list_item_prefix(node: object, kind: str, index: int) -> str:
    task = getattr(node, "task", False)
    if task:
        checked = getattr(node, "checked", None)
        return "\u2611" if checked is True else "\u2610"
    if kind == "ordered":
        return f"{index}."
    return "-"


def _list_item_indent(node: object) -> str:
    level = getattr(node, "level", 0)
    return "  " * max(level, 0)


def _coerce_cell_inlines(value: object) -> list[object]:
    if isinstance(value, str):
        return [type("InlineText", (), {"text": value})()]
    return list(value)


def _render_table_cell(
    cell,
    value: object,
    *,
    style_name: str | None = None,
    properties: Mapping[str, object] | None = None,
) -> None:
    paragraph = cell.paragraphs[0]
    if style_name is not None:
        paragraph.style = style_name
    if properties is not None:
        _apply_paragraph_properties(paragraph, properties)
    if paragraph.runs:
        for run in list(paragraph.runs):
            run._element.getparent().remove(run._element)
    _add_inline_runs(paragraph, _coerce_cell_inlines(value), properties=properties)


def _apply_paragraph_properties(paragraph, properties: Mapping[str, object]) -> None:
    paragraph_format = paragraph.paragraph_format
    for key, attr_name in (
        ("first_line_indent_pt", "first_line_indent"),
        ("left_indent_pt", "left_indent"),
        ("right_indent_pt", "right_indent"),
        ("space_before_pt", "space_before"),
        ("space_after_pt", "space_after"),
    ):
        value = properties.get(key)
        if isinstance(value, (int, float)):
            setattr(paragraph_format, attr_name, Pt(float(value)))
    line_spacing = properties.get("line_spacing")
    line_spacing_pt = properties.get("line_spacing_pt")
    if isinstance(line_spacing, (int, float)):
        paragraph_format.line_spacing = float(line_spacing)
    elif isinstance(line_spacing_pt, (int, float)):
        paragraph_format.line_spacing = Pt(float(line_spacing_pt))


def _style_numbering_reference(style) -> tuple[int, int] | None:
    current = style
    while current is not None:
        properties = getattr(current.element, "pPr", None)
        num_pr = getattr(properties, "numPr", None)
        if num_pr is not None and getattr(num_pr, "numId", None) is not None:
            ilvl = getattr(num_pr, "ilvl", None)
            return int(num_pr.numId.val), int(ilvl.val if ilvl is not None else 0)
        current = getattr(current, "base_style", None)
    return None


def _build_style_numbering_map(document: DocumentType) -> dict[str, tuple[int, int]]:
    mapping: dict[str, tuple[int, int]] = {}
    for style in document.styles:
        name = getattr(style, "name", "")
        if not name:
            continue
        reference = _style_numbering_reference(style)
        if reference is not None:
            mapping[name] = reference
    return mapping


def _create_numbering_instance(
    document: DocumentType,
    *,
    base_num_id: int,
    ilvl: int,
    start_at: int,
) -> int | None:
    try:
        numbering = document.part.numbering_part.element
    except KeyError:
        return None

    base_nodes = numbering.xpath(f"./w:num[@w:numId='{base_num_id}']")
    if not base_nodes:
        return None
    base_node = base_nodes[0]
    abstract_refs = base_node.xpath("./w:abstractNumId/@w:val")
    if not abstract_refs:
        return None
    abstract_id = int(abstract_refs[0])

    existing_ids = [int(value) for value in numbering.xpath("./w:num/@w:numId")]
    next_num_id = max(existing_ids, default=0) + 1

    num_node = OxmlElement("w:num")
    num_node.set(qn("w:numId"), str(next_num_id))

    abstract_node = OxmlElement("w:abstractNumId")
    abstract_node.set(qn("w:val"), str(abstract_id))
    num_node.append(abstract_node)

    if start_at > 1:
        level_override = OxmlElement("w:lvlOverride")
        level_override.set(qn("w:ilvl"), str(max(ilvl, 0)))
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), str(start_at))
        level_override.append(start_override)
        num_node.append(level_override)

    numbering.append(num_node)
    return next_num_id


def _bind_numbering(
    paragraph,
    style_numbering_map: Mapping[str, tuple[int, int]],
    numbering: NumberingIntent | None,
    *,
    document: DocumentType | None = None,
    numbering_state: dict[tuple[str, str, int], int] | None = None,
) -> bool:
    numbering_style = numbering.preferred_template_style if numbering is not None else None
    if not numbering_style:
        return False
    reference = style_numbering_map.get(numbering_style)
    if reference is None:
        return False
    num_id, template_ilvl = reference
    ilvl = max(numbering.level, 0) if numbering is not None else template_ilvl
    if numbering is not None and numbering_state is not None:
        key = (numbering.numbering_role, numbering_style, ilvl)
        if numbering.continue_sequence:
            num_id = numbering_state.get(key, num_id)
            numbering_state.setdefault(key, num_id)
        else:
            start_at = numbering.start_at if isinstance(numbering.start_at, int) and numbering.start_at > 0 else 1
            if document is not None:
                restarted_num_id = _create_numbering_instance(
                    document,
                    base_num_id=num_id,
                    ilvl=ilvl,
                    start_at=start_at,
                )
                if restarted_num_id is not None:
                    num_id = restarted_num_id
            numbering_state[key] = num_id
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    num_pr.get_or_add_numId().val = num_id
    num_pr.get_or_add_ilvl().val = ilvl
    return True


def _render_single_layout_intent(
    document: DocumentType,
    intent: ParagraphLayoutIntent,
    style_numbering_map: Mapping[str, tuple[int, int]],
    numbering_state: dict[tuple[str, str, int], int],
) -> None:
    paragraph = document.add_paragraph(style=intent.resolved_style_name)
    _apply_paragraph_properties(paragraph, intent.resolved_properties)
    numbered = _bind_numbering(
        paragraph,
        style_numbering_map,
        intent.numbering,
        document=document,
        numbering_state=numbering_state,
    )
    if not numbered and intent.prefix_text:
        paragraph.add_run(intent.prefix_text)
    _add_inline_runs(paragraph, intent.runs, properties=intent.resolved_properties)


def _render_layout_intents(
    document: DocumentType,
    layout_intents: list[ParagraphLayoutIntent],
    style_numbering_map: Mapping[str, tuple[int, int]],
) -> None:
    numbering_state: dict[tuple[str, str, int], int] = {}
    for intent in layout_intents:
        _render_single_layout_intent(document, intent, style_numbering_map, numbering_state)


def _render_horizontal_rule_node(document: DocumentType, styles: Mapping[str, object]) -> None:
    style_name = styles["paragraph"].word_style_name
    document.add_paragraph("---", style=style_name)


def _render_table_node(document: DocumentType, node: TableNode, styles: Mapping[str, object]) -> None:
    total_rows = 1 + len(node.rows)
    total_cols = max(len(node.headers), max((len(row) for row in node.rows), default=0))
    table = document.add_table(rows=total_rows, cols=total_cols)
    table_style = styles.get("table", styles["paragraph"]) if "table" in styles else None
    table_style_name = table_style.word_style_name if table_style is not None else None
    table_properties = getattr(table_style, "resolved_properties", {}) if table_style is not None else {}
    if table_style_name is not None:
        table.style = table_style_name
    for col_index, value in enumerate(node.headers):
        _render_table_cell(table.cell(0, col_index), value, properties=table_properties)
    for row_index, row in enumerate(node.rows, start=1):
        for col_index, value in enumerate(row):
            _render_table_cell(table.cell(row_index, col_index), value, properties=table_properties)


def _render_code_block_node(document: DocumentType, node: CodeBlockNode, styles: Mapping[str, object]) -> None:
    table = document.add_table(rows=1, cols=1)
    code_style = styles.get("code_block", styles["paragraph"])
    code_style_name = code_style.word_style_name
    code_properties = getattr(code_style, "resolved_properties", {})
    _render_table_cell(table.cell(0, 0), node.content, style_name=code_style_name, properties=code_properties)


def _render_image_node(
    document: DocumentType,
    node: ImageBlockNode,
    styles: Mapping[str, object],
    base_dir: Path | None,
) -> None:
    try:
        if base_dir is None:
            raise ValueError("Image rendering requires a base_dir")
        image_path = resolve_image_path(node.raw_path, base_dir)
        if not image_path.exists():
            raise FileNotFoundError(str(image_path))
        document.add_picture(str(image_path))
    except (OSError, ValueError):
        style_name = styles.get("image", styles["paragraph"]).word_style_name
        document.add_paragraph(
            render_missing_image_placeholder(node.raw_path),
            style=style_name,
        )


def render_nodes_to_document(
    nodes: Iterable[object],
    styles: Mapping[str, object],
    *,
    base_dir: Path | None = None,
    template_path: Path | None = None,
    layout_intents: list[ParagraphLayoutIntent] | None = None,
) -> DocumentType:
    document = create_document(template_path)
    style_numbering_map = _build_style_numbering_map(document)
    nodes = list(nodes)

    if layout_intents is not None:
        numbering_state: dict[tuple[str, str, int], int] = {}
        intent_index = 0
        for node in nodes:
            if isinstance(node, HorizontalRuleNode):
                _render_horizontal_rule_node(document, styles)
            elif isinstance(node, TableNode):
                _render_table_node(document, node, styles)
            elif isinstance(node, CodeBlockNode):
                _render_code_block_node(document, node, styles)
            elif isinstance(node, ImageBlockNode):
                _render_image_node(document, node, styles, base_dir)
            elif isinstance(node, ListNode):
                item_count = len(node.items)
                for intent in layout_intents[intent_index:intent_index + item_count]:
                    _render_single_layout_intent(document, intent, style_numbering_map, numbering_state)
                intent_index += item_count
            else:
                intent = layout_intents[intent_index]
                _render_single_layout_intent(document, intent, style_numbering_map, numbering_state)
                intent_index += 1
        while intent_index < len(layout_intents):
            _render_single_layout_intent(document, layout_intents[intent_index], style_numbering_map, numbering_state)
            intent_index += 1
        return document

    numbering_state: dict[tuple[str, str, int], int] = {}
    for node in nodes:
        if isinstance(node, HeadingNode):
            style_name = styles[f"heading{node.level}"].word_style_name
            paragraph = document.add_paragraph(style=style_name)
            _apply_paragraph_properties(paragraph, getattr(styles[f"heading{node.level}"], "resolved_properties", {}))
            _add_inline_runs(paragraph, node.inlines, properties=getattr(styles[f"heading{node.level}"], "resolved_properties", {}))
            continue

        if isinstance(node, ParagraphNode):
            style_name = styles["paragraph"].word_style_name
            paragraph = document.add_paragraph(style=style_name)
            _apply_paragraph_properties(paragraph, getattr(styles["paragraph"], "resolved_properties", {}))
            _add_inline_runs(paragraph, node.inlines, properties=getattr(styles["paragraph"], "resolved_properties", {}))
            continue

        if isinstance(node, QuoteNode):
            style_name = styles.get("quote", styles["paragraph"]).word_style_name
            paragraph = document.add_paragraph(style=style_name)
            _apply_paragraph_properties(paragraph, getattr(styles.get("quote", styles["paragraph"]), "resolved_properties", {}))
            paragraph.add_run("> ")
            _add_inline_runs(paragraph, node.inlines, properties=getattr(styles.get("quote", styles["paragraph"]), "resolved_properties", {}))
            continue

        if isinstance(node, HorizontalRuleNode):
            _render_horizontal_rule_node(document, styles)
            continue

        if isinstance(node, ListNode):
            ordered_counters: dict[int, int] = {}
            raw_start = getattr(node, "start", 1)
            try:
                ordered_list_start = max(int(raw_start), 1)
            except (TypeError, ValueError):
                ordered_list_start = 1
            first_ordered_item = True
            for item in node.items:
                item_kind = getattr(item, "kind", None) or node.kind
                if item_kind == "ordered":
                    ordered_style = styles.get("ordered_list", styles["paragraph"])
                    style_name = ordered_style.word_style_name
                    resolved_properties = getattr(ordered_style, "resolved_properties", {})
                else:
                    unordered_style = styles.get("unordered_list", styles["paragraph"])
                    style_name = unordered_style.word_style_name
                    resolved_properties = getattr(unordered_style, "resolved_properties", {})
                item_level = max(getattr(item, "level", 0), 0)
                raw_sequence_start = getattr(item, "sequence_start", None)
                try:
                    sequence_start = max(int(raw_sequence_start), 1) if raw_sequence_start is not None else None
                except (TypeError, ValueError):
                    sequence_start = None
                if sequence_start is None and first_ordered_item:
                    sequence_start = ordered_list_start
                ordered_counters = {
                    level: counter for level, counter in ordered_counters.items() if level <= item_level
                }
                display_index = 1
                if item_kind == "ordered":
                    if sequence_start is not None:
                        ordered_counters[item_level] = sequence_start - 1
                    elif item_level == 0 and 0 not in ordered_counters:
                        ordered_counters[0] = ordered_list_start - 1
                    display_index = ordered_counters.get(item_level, 0) + 1
                    ordered_counters[item_level] = display_index
                paragraph = document.add_paragraph(style=style_name)
                _apply_paragraph_properties(paragraph, resolved_properties)
                numbered = False
                if item_kind == "ordered":
                    numbered = _bind_numbering(
                        paragraph,
                        style_numbering_map,
                        NumberingIntent(
                            numbering_role="ordered_list",
                            level=item_level,
                            continue_sequence=sequence_start is None,
                            start_at=sequence_start,
                            preferred_template_style=(
                                resolved_properties.get("numbering_style")
                                if isinstance(resolved_properties, Mapping)
                                else None
                            ),
                        ),
                        document=document,
                        numbering_state=numbering_state,
                    )
                    if first_ordered_item:
                        first_ordered_item = False
                elif item_kind == "unordered" and not getattr(item, "task", False):
                    numbered = _bind_numbering(
                        paragraph,
                        style_numbering_map,
                        NumberingIntent(
                            numbering_role="unordered_list",
                            level=item_level,
                            preferred_template_style=(
                                resolved_properties.get("numbering_style")
                                if isinstance(resolved_properties, Mapping)
                                else None
                            ),
                        ),
                        document=document,
                        numbering_state=numbering_state,
                    )
                if not numbered:
                    paragraph.add_run(
                        f"{_list_item_indent(item)}{_list_item_prefix(item, item_kind, display_index)} "
                    )
                _add_inline_runs(paragraph, item.inlines, properties=resolved_properties)
            continue

        if isinstance(node, TableNode):
            _render_table_node(document, node, styles)
            continue

        if isinstance(node, CodeBlockNode):
            _render_code_block_node(document, node, styles)
            continue

        if isinstance(node, ImageBlockNode):
            _render_image_node(document, node, styles, base_dir)
            continue

        raise TypeError(f"Unsupported node type: {type(node).__name__}")

    return document
