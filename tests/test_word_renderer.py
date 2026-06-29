from pathlib import Path
from types import SimpleNamespace
import base64

import pytest
from docx import Document
from docx import Document as WordDocument
from docx.oxml.ns import qn

from docubridge.application.models import RenderRequest
from docubridge.application.render_service import run_render
from docubridge.core.asset_resolver import resolve_image_path
from docubridge.core.layout_intent import NumberingIntent, ParagraphLayoutIntent, RunStyleIntent
from docubridge.core.nodes import CodeBlockNode, HeadingNode, HorizontalRuleNode, ImageBlockNode, ListItemNode, ListNode, ParagraphNode, QuoteNode, TableNode, TextSpan
from docubridge.core.word_renderer import render_missing_image_placeholder, render_nodes_to_document

_PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+aRX0AAAAASUVORK5CYII="
)


def _paragraph_numbering_reference(paragraph) -> tuple[int, int] | None:
    num_pr = paragraph._p.pPr.numPr if paragraph._p.pPr is not None else None
    if num_pr is None or num_pr.numId is None or num_pr.ilvl is None:
        return None
    return int(num_pr.numId.val), int(num_pr.ilvl.val)


def _numbering_start_override(document: Document, *, num_id: int, ilvl: int) -> int | None:
    numbering = document.part.numbering_part.element
    num_nodes = numbering.xpath(f"./w:num[@w:numId='{num_id}']")
    if not num_nodes:
        return None
    override_nodes = num_nodes[0].xpath(f"./w:lvlOverride[@w:ilvl='{ilvl}']")
    if not override_nodes:
        return None
    start_nodes = override_nodes[0].xpath("./w:startOverride/@w:val")
    if not start_nodes:
        return None
    return int(start_nodes[0])


def test_render_nodes_to_document_creates_heading_and_paragraph() -> None:
    nodes = [
        HeadingNode(
            level=1,
            inlines=[TextSpan(text="Sample Title")],
        ),
        ParagraphNode(
            inlines=[TextSpan(text="Plain paragraph")],
        ),
    ]
    styles = {
        "heading1": SimpleNamespace(word_style_name="Heading 1"),
        "paragraph": SimpleNamespace(word_style_name="Normal"),
    }

    document = render_nodes_to_document(nodes, styles)

    assert len(document.paragraphs) == 2
    assert document.paragraphs[0].text == "Sample Title"
    assert document.paragraphs[0].style.name == "Heading 1"
    assert document.paragraphs[1].text == "Plain paragraph"
    assert document.paragraphs[1].style.name == "Normal"


def test_render_nodes_to_document_preserves_inline_runs_inside_headings() -> None:
    nodes = [
        HeadingNode(
            level=1,
            inlines=[
                TextSpan(text="Title "),
                TextSpan(text="bold", bold=True),
                TextSpan(text=" "),
                TextSpan(text="italic", italic=True),
                TextSpan(text=" "),
                TextSpan(text="gone", strike=True),
                TextSpan(text=" "),
                TextSpan(text="code", code=True),
                TextSpan(text=" "),
                TextSpan(text="OpenAI", href="https://openai.com"),
            ],
        )
    ]
    styles = {
        "heading1": SimpleNamespace(word_style_name="Heading 1"),
        "paragraph": SimpleNamespace(word_style_name="Normal"),
    }

    document = render_nodes_to_document(nodes, styles)

    paragraph = document.paragraphs[0]
    assert paragraph.text == "Title bold italic gone code OpenAI"
    assert paragraph.style.name == "Heading 1"
    assert any(run.bold for run in paragraph.runs if run.text == "bold")
    assert any(run.italic for run in paragraph.runs if run.text == "italic")
    assert any(run.font.strike for run in paragraph.runs if run.text == "gone")
    assert any(run.font.name == "Consolas" for run in paragraph.runs if run.text == "code")
    assert "hyperlink" in paragraph._p.xml


def test_render_nodes_to_document_preserves_inline_bold_italic_and_code_runs() -> None:
    nodes = [
        ParagraphNode(
            inlines=[
                TextSpan(text="Plain "),
                TextSpan(text="bold", bold=True),
                TextSpan(text=" "),
                TextSpan(text="italic", italic=True),
                TextSpan(text=" "),
                TextSpan(text="code", code=True),
            ],
        )
    ]
    styles = {"paragraph": SimpleNamespace(word_style_name="Normal")}

    document = render_nodes_to_document(nodes, styles)

    runs = document.paragraphs[0].runs
    assert [run.text for run in runs] == ["Plain ", "bold", " ", "italic", " ", "code"]
    assert runs[1].bold is True
    assert runs[3].italic is True
    assert runs[5].font.name == "Consolas"


def test_render_nodes_to_document_applies_font_size_from_resolved_properties() -> None:
    nodes = [
        ParagraphNode(inlines=[TextSpan(text="Body text")]),
    ]
    styles = {
        "paragraph": SimpleNamespace(
            word_style_name="Normal",
            resolved_properties={"font_size": 14},
        ),
    }

    document = render_nodes_to_document(nodes, styles)

    run = document.paragraphs[0].runs[0]
    assert run.font.size is not None
    assert run.font.size.pt == 14


def test_render_nodes_to_document_creates_hyperlinks_for_inline_links() -> None:
    nodes = [
        ParagraphNode(
            inlines=[
                TextSpan(text="Visit "),
                TextSpan(text="OpenAI", href="https://openai.com"),
            ],
        )
    ]
    styles = {"paragraph": SimpleNamespace(word_style_name="Normal")}

    document = render_nodes_to_document(nodes, styles)
    xml = document.paragraphs[0]._p.xml

    assert "hyperlink" in xml
    assert "OpenAI" in document.paragraphs[0].text


def test_render_nodes_to_document_preserves_inline_strikethrough_runs() -> None:
    nodes = [
        ParagraphNode(
            inlines=[
                TextSpan(text="Keep "),
                TextSpan(text="remove", strike=True),
                TextSpan(text=" text"),
            ],
        )
    ]
    styles = {"paragraph": SimpleNamespace(word_style_name="Normal")}

    document = render_nodes_to_document(nodes, styles)

    runs = document.paragraphs[0].runs
    assert [run.text for run in runs] == ["Keep ", "remove", " text"]
    assert runs[1].font.strike is True


def test_render_nodes_to_document_formats_list_nodes() -> None:
    nodes = [
        ListNode(
            kind="ordered",
            items=[
                ListItemNode(inlines=[TextSpan(text="First")]),
                ListItemNode(inlines=[TextSpan(text="Second")]),
            ],
        ),
        ListNode(
            kind="unordered",
            items=[
                ListItemNode(inlines=[TextSpan(text="Bullet")]),
            ],
        ),
        ListNode(
            kind="unordered",
            items=[
                ListItemNode(inlines=[TextSpan(text="Done")], task=True, checked=True),
                ListItemNode(inlines=[TextSpan(text="Todo")], task=True, checked=False),
            ],
        ),
    ]
    styles = {"paragraph": SimpleNamespace(word_style_name="Normal")}

    document = render_nodes_to_document(nodes, styles)

    assert [paragraph.text for paragraph in document.paragraphs] == [
        "1. First",
        "2. Second",
        "- Bullet",
        "☑ Done",
        "☐ Todo",
    ]


def test_render_nodes_to_document_uses_list_and_quote_styles_when_available() -> None:
    nodes = [
        ListNode(
            kind="ordered",
            items=[ListItemNode(inlines=[TextSpan(text="First")], kind="ordered")],
        ),
        ListNode(
            kind="unordered",
            items=[ListItemNode(inlines=[TextSpan(text="Bullet")], kind="unordered")],
        ),
        QuoteNode(inlines=[TextSpan(text="Quoted")]),
    ]
    styles = {
        "paragraph": SimpleNamespace(word_style_name="Normal"),
        "ordered_list": SimpleNamespace(word_style_name="List Number"),
        "unordered_list": SimpleNamespace(word_style_name="List Bullet"),
        "quote": SimpleNamespace(word_style_name="Quote"),
    }

    document = render_nodes_to_document(nodes, styles)

    assert document.paragraphs[0].style.name == "List Number"
    assert document.paragraphs[1].style.name == "List Bullet"
    assert document.paragraphs[2].style.name == "Quote"


def test_render_nodes_to_document_preserves_nested_list_indentation() -> None:
    nodes = [
        ListNode(
            kind="unordered",
            items=[
                ListItemNode(inlines=[TextSpan(text="Parent")], level=0, kind="unordered"),
                ListItemNode(inlines=[TextSpan(text="Child")], level=1, kind="unordered"),
            ],
        ),
        ListNode(
            kind="ordered",
            items=[
                ListItemNode(inlines=[TextSpan(text="Ordered parent")], level=0, kind="ordered"),
                ListItemNode(inlines=[TextSpan(text="Ordered child")], level=1, kind="ordered"),
            ],
        ),
    ]
    styles = {"paragraph": SimpleNamespace(word_style_name="Normal")}

    document = render_nodes_to_document(nodes, styles)

    assert [paragraph.text for paragraph in document.paragraphs] == [
        "- Parent",
        "  - Child",
        "1. Ordered parent",
        "  1. Ordered child",
    ]


def test_render_nodes_to_document_preserves_inline_runs_inside_list_items() -> None:
    nodes = [
        ListNode(
            kind="unordered",
            items=[
                ListItemNode(
                    inlines=[
                        TextSpan(text="Plain "),
                        TextSpan(text="bold", bold=True),
                        TextSpan(text=" "),
                        TextSpan(text="italic", italic=True),
                        TextSpan(text=" "),
                        TextSpan(text="code", code=True),
                        TextSpan(text=" "),
                        TextSpan(text="OpenAI", href="https://openai.com"),
                    ]
                )
            ],
        )
    ]
    styles = {"paragraph": SimpleNamespace(word_style_name="Normal")}

    document = render_nodes_to_document(nodes, styles)

    paragraph = document.paragraphs[0]
    assert paragraph.text == "- Plain bold italic code OpenAI"
    assert any(run.bold for run in paragraph.runs if run.text == "bold")
    assert any(run.italic for run in paragraph.runs if run.text == "italic")
    assert any(run.font.name == "Consolas" for run in paragraph.runs if run.text == "code")
    assert "hyperlink" in paragraph._p.xml


def test_render_nodes_to_document_binds_native_numbering_and_paragraph_properties(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    document = Document()
    document.save(template_path)
    nodes = [
        ListNode(
            kind="ordered",
            items=[ListItemNode(inlines=[TextSpan(text="Alpha")], level=0, kind="ordered")],
        )
    ]
    styles = {
        "paragraph": SimpleNamespace(word_style_name="Normal", resolved_properties={}),
        "ordered_list": SimpleNamespace(
            word_style_name="List Number",
            resolved_properties={
                "first_line_indent_pt": 21,
                "left_indent_pt": 28,
                "numbering_style": "List Number",
            },
        ),
    }

    rendered = render_nodes_to_document(nodes, styles, template_path=template_path)

    paragraph = rendered.paragraphs[0]
    assert paragraph.style.name == "List Number"
    assert _paragraph_numbering_reference(paragraph) is not None
    assert paragraph.text == "Alpha"


def test_render_nodes_to_document_uses_layout_intents_for_numbering_and_run_properties(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    document = Document()
    document.save(template_path)
    nodes = [
        ListNode(
            kind="ordered",
            items=[ListItemNode(inlines=[TextSpan(text="Alpha")], level=0, kind="ordered")],
        )
    ]
    styles = {
        "paragraph": SimpleNamespace(word_style_name="Normal", resolved_properties={}),
        "ordered_list": SimpleNamespace(
            word_style_name="List Number",
            resolved_properties={},
        ),
    }
    layout_intents = [
        ParagraphLayoutIntent(
            element_name="ordered_list",
            runs=[RunStyleIntent(text="Alpha")],
            resolved_style_name="List Number",
            resolved_properties={"font_ascii": "Times New Roman"},
            numbering=NumberingIntent(
                numbering_role="ordered_list",
                level=0,
                preferred_template_style="List Number",
            ),
        )
    ]

    rendered = render_nodes_to_document(
        nodes,
        styles,
        template_path=template_path,
        layout_intents=layout_intents,
    )

    paragraph = rendered.paragraphs[0]
    assert paragraph.style.name == "List Number"
    assert _paragraph_numbering_reference(paragraph) is not None
    assert paragraph.text == "Alpha"
    assert paragraph.runs[0].font.name == "Times New Roman"


def test_render_nodes_to_document_uses_native_bullet_numbering_from_layout_intents(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    document = Document()
    document.save(template_path)
    nodes = [
        ListNode(
            kind="unordered",
            items=[ListItemNode(inlines=[TextSpan(text="Bullet")], level=0, kind="unordered")],
        )
    ]
    styles = {
        "paragraph": SimpleNamespace(word_style_name="Normal", resolved_properties={}),
        "unordered_list": SimpleNamespace(
            word_style_name="List Bullet",
            resolved_properties={"numbering_style": "List Bullet"},
        ),
    }
    layout_intents = [
        ParagraphLayoutIntent(
            element_name="unordered_list",
            runs=[RunStyleIntent(text="Bullet")],
            resolved_style_name="List Bullet",
            resolved_properties={},
            prefix_text="- ",
            numbering=NumberingIntent(
                numbering_role="unordered_list",
                level=0,
                preferred_template_style="List Bullet",
            ),
        )
    ]

    rendered = render_nodes_to_document(
        nodes,
        styles,
        template_path=template_path,
        layout_intents=layout_intents,
    )

    paragraph = rendered.paragraphs[0]
    assert paragraph.style.name == "List Bullet"
    assert _paragraph_numbering_reference(paragraph) is not None
    assert paragraph.text == "Bullet"


def test_render_nodes_to_document_applies_nested_numbering_level_from_layout_intents(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    document = Document()
    document.save(template_path)
    nodes = [
        ListNode(
            kind="ordered",
            items=[ListItemNode(inlines=[TextSpan(text="Nested")], level=1, kind="ordered")],
        )
    ]
    styles = {
        "paragraph": SimpleNamespace(word_style_name="Normal", resolved_properties={}),
        "ordered_list": SimpleNamespace(
            word_style_name="List Number",
            resolved_properties={},
        ),
    }
    layout_intents = [
        ParagraphLayoutIntent(
            element_name="ordered_list",
            runs=[RunStyleIntent(text="Nested")],
            resolved_style_name="List Number",
            resolved_properties={},
            prefix_text="  1. ",
            numbering=NumberingIntent(
                numbering_role="ordered_list",
                level=1,
                preferred_template_style="List Number",
            ),
        )
    ]

    rendered = render_nodes_to_document(
        nodes,
        styles,
        template_path=template_path,
        layout_intents=layout_intents,
    )

    paragraph = rendered.paragraphs[0]
    assert paragraph.text == "Nested"
    assert _paragraph_numbering_reference(paragraph) is not None
    assert _paragraph_numbering_reference(paragraph)[1] == 1


def test_render_nodes_to_document_binds_nested_unordered_list_to_nested_numbering_level(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    document = Document()
    document.save(template_path)
    nodes = [
        ListNode(
            kind="unordered",
            items=[
                ListItemNode(inlines=[TextSpan(text="Parent")], level=0, kind="unordered"),
                ListItemNode(inlines=[TextSpan(text="Child")], level=1, kind="unordered"),
            ],
        )
    ]
    styles = {
        "paragraph": SimpleNamespace(word_style_name="Normal", resolved_properties={}),
        "unordered_list": SimpleNamespace(
            word_style_name="List Bullet",
            resolved_properties={"numbering_style": "List Bullet"},
        ),
    }

    rendered = render_nodes_to_document(
        nodes,
        styles,
        template_path=template_path,
    )

    assert [paragraph.text for paragraph in rendered.paragraphs] == ["Parent", "Child"]
    assert _paragraph_numbering_reference(rendered.paragraphs[0]) is not None
    assert _paragraph_numbering_reference(rendered.paragraphs[1]) is not None
    assert _paragraph_numbering_reference(rendered.paragraphs[0])[1] == 0
    assert _paragraph_numbering_reference(rendered.paragraphs[1])[1] == 1


def test_render_nodes_to_document_restarts_ordered_sequence_with_start_override_from_layout_intents(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    document = Document()
    document.save(template_path)
    nodes = []
    styles = {
        "paragraph": SimpleNamespace(word_style_name="Normal", resolved_properties={}),
        "ordered_list": SimpleNamespace(
            word_style_name="List Number",
            resolved_properties={},
        ),
    }
    layout_intents = [
        ParagraphLayoutIntent(
            element_name="ordered_list",
            runs=[RunStyleIntent(text="First")],
            resolved_style_name="List Number",
            resolved_properties={},
            numbering=NumberingIntent(
                numbering_role="ordered_list",
                level=0,
                continue_sequence=False,
                start_at=1,
                preferred_template_style="List Number",
            ),
        ),
        ParagraphLayoutIntent(
            element_name="ordered_list",
            runs=[RunStyleIntent(text="Second")],
            resolved_style_name="List Number",
            resolved_properties={},
            numbering=NumberingIntent(
                numbering_role="ordered_list",
                level=0,
                continue_sequence=True,
                preferred_template_style="List Number",
            ),
        ),
        ParagraphLayoutIntent(
            element_name="ordered_list",
            runs=[RunStyleIntent(text="Third")],
            resolved_style_name="List Number",
            resolved_properties={},
            numbering=NumberingIntent(
                numbering_role="ordered_list",
                level=0,
                continue_sequence=False,
                start_at=3,
                preferred_template_style="List Number",
            ),
        ),
        ParagraphLayoutIntent(
            element_name="ordered_list",
            runs=[RunStyleIntent(text="Fourth")],
            resolved_style_name="List Number",
            resolved_properties={},
            numbering=NumberingIntent(
                numbering_role="ordered_list",
                level=0,
                continue_sequence=True,
                preferred_template_style="List Number",
            ),
        ),
    ]

    rendered = render_nodes_to_document(
        nodes,
        styles,
        template_path=template_path,
        layout_intents=layout_intents,
    )

    first_ref = _paragraph_numbering_reference(rendered.paragraphs[0])
    second_ref = _paragraph_numbering_reference(rendered.paragraphs[1])
    third_ref = _paragraph_numbering_reference(rendered.paragraphs[2])
    fourth_ref = _paragraph_numbering_reference(rendered.paragraphs[3])
    assert first_ref is not None and second_ref is not None and third_ref is not None and fourth_ref is not None
    assert first_ref[0] == second_ref[0]
    assert third_ref[0] == fourth_ref[0]
    assert third_ref[0] != first_ref[0]
    assert _numbering_start_override(rendered, num_id=third_ref[0], ilvl=0) == 3


def test_render_nodes_to_document_preserves_task_list_inline_runs() -> None:
    nodes = [
        ListNode(
            kind="unordered",
            items=[
                ListItemNode(
                    inlines=[
                        TextSpan(text="done", bold=True),
                        TextSpan(text=" "),
                        TextSpan(text="old", strike=True),
                        TextSpan(text=" "),
                        TextSpan(text="code", code=True),
                        TextSpan(text=" "),
                        TextSpan(text="ref", href="https://example.com"),
                    ],
                    task=True,
                    checked=True,
                )
            ],
        )
    ]
    styles = {"paragraph": SimpleNamespace(word_style_name="Normal")}

    document = render_nodes_to_document(nodes, styles)

    paragraph = document.paragraphs[0]
    assert paragraph.text == "☑ done old code ref"
    assert any(run.bold for run in paragraph.runs if run.text == "done")
    assert any(run.font.strike for run in paragraph.runs if run.text == "old")
    assert any(run.font.name == "Consolas" for run in paragraph.runs if run.text == "code")
    assert "hyperlink" in paragraph._p.xml


def test_render_nodes_to_document_creates_docx_tables() -> None:
    nodes = [
        TableNode(
            headers=[
                [TextSpan(text="Name")],
                [TextSpan(text="Value")],
            ],
            rows=[
                [
                    [TextSpan(text="Alpha")],
                    [TextSpan(text="1")],
                ],
                [
                    [TextSpan(text="Beta")],
                    [TextSpan(text="2")],
                ],
            ],
        )
    ]
    styles = {"paragraph": SimpleNamespace(word_style_name="Normal")}

    document = render_nodes_to_document(nodes, styles)

    assert len(document.tables) == 1
    table = document.tables[0]
    assert table.cell(0, 0).text == "Name"
    assert table.cell(0, 1).text == "Value"
    assert table.cell(1, 0).text == "Alpha"
    assert table.cell(1, 1).text == "1"
    assert table.cell(2, 0).text == "Beta"
    assert table.cell(2, 1).text == "2"


def test_render_nodes_to_document_applies_heading_font_slots_to_runs() -> None:
    nodes = [
        HeadingNode(
            level=1,
            inlines=[TextSpan(text="标题")],
        )
    ]
    styles = {
        "heading1": SimpleNamespace(
            word_style_name="Heading 1",
            resolved_properties={
                "font_ascii": "Times New Roman",
                "font_east_asia": "宋体",
            },
        ),
        "paragraph": SimpleNamespace(word_style_name="Normal"),
    }

    document = render_nodes_to_document(nodes, styles)

    run = document.paragraphs[0].runs[0]
    assert run.font.name == "Times New Roman"
    assert 'w:eastAsia="宋体"' in run._element.xml


def test_render_nodes_to_document_preserves_inline_runs_inside_table_cells() -> None:
    nodes = [
        TableNode(
            headers=[
                [TextSpan(text="Name")],
                [TextSpan(text="Value")],
            ],
            rows=[
                [
                    [TextSpan(text="Alpha", bold=True)],
                    [
                        TextSpan(text="42", code=True),
                        TextSpan(text=" "),
                        TextSpan(text="ref", href="https://example.com"),
                    ],
                ]
            ],
        )
    ]
    styles = {"paragraph": SimpleNamespace(word_style_name="Normal")}

    document = render_nodes_to_document(nodes, styles)

    table = document.tables[0]
    first_cell_runs = table.cell(1, 0).paragraphs[0].runs
    second_cell_paragraph = table.cell(1, 1).paragraphs[0]
    assert table.cell(1, 0).text == "Alpha"
    assert table.cell(1, 1).text == "42 ref"
    assert any(run.bold for run in first_cell_runs if run.text == "Alpha")
    assert any(run.font.name == "Consolas" for run in second_cell_paragraph.runs if run.text == "42")
    assert "hyperlink" in second_cell_paragraph._p.xml


def test_render_nodes_to_document_applies_table_cell_paragraph_properties() -> None:
    nodes = [
        TableNode(
            headers=[[TextSpan(text="Name")]],
            rows=[[[TextSpan(text="Alpha")]]],
        )
    ]
    styles = {
        "paragraph": SimpleNamespace(word_style_name="Normal"),
        "table": SimpleNamespace(
            word_style_name="Table Grid",
            resolved_properties={"space_after_pt": 6},
        ),
    }

    document = render_nodes_to_document(nodes, styles)

    paragraph = document.tables[0].cell(1, 0).paragraphs[0]
    assert paragraph.paragraph_format.space_after.pt == 6


def test_render_nodes_to_document_creates_code_block_tables() -> None:
    nodes = [CodeBlockNode(content="print('hello')\nprint('world')\n", language="python")]
    styles = {"paragraph": SimpleNamespace(word_style_name="Normal")}

    document = render_nodes_to_document(nodes, styles)

    assert len(document.tables) == 1
    assert document.tables[0].cell(0, 0).text == "print('hello')\nprint('world')\n"


def test_render_nodes_to_document_applies_code_block_font_slots() -> None:
    nodes = [CodeBlockNode(content="print('hello')\n", language="python")]
    styles = {
        "paragraph": SimpleNamespace(word_style_name="Normal"),
        "code_block": SimpleNamespace(
            word_style_name="Quote",
            resolved_properties={
                "font_ascii": "Courier New",
                "font_east_asia": "等线",
            },
        ),
    }

    document = render_nodes_to_document(nodes, styles)

    paragraph = document.tables[0].cell(0, 0).paragraphs[0]
    run = paragraph.runs[0]
    assert run.font.name == "Courier New"
    assert 'w:eastAsia="等线"' in run._element.xml


def test_render_nodes_to_document_uses_table_and_code_block_styles_when_available() -> None:
    nodes = [
        TableNode(
            headers=[[TextSpan(text="Name")]],
            rows=[[[TextSpan(text="Alpha")]]],
        ),
        CodeBlockNode(content="print('hello')\n", language="python"),
    ]
    styles = {
        "paragraph": SimpleNamespace(word_style_name="Normal"),
        "table": SimpleNamespace(word_style_name="Table Grid"),
        "code_block": SimpleNamespace(word_style_name="Quote"),
    }

    document = render_nodes_to_document(nodes, styles)

    assert document.tables[0].style.name == "Table Grid"
    assert document.tables[1].cell(0, 0).paragraphs[0].style.name == "Quote"


def test_render_nodes_to_document_inserts_images(tmp_path: Path) -> None:
    image_path = tmp_path / "diagram.png"
    image_path.write_bytes(_PNG_1X1)
    nodes = [ImageBlockNode(raw_path=image_path.name, alt_text="diagram")]
    styles = {"paragraph": SimpleNamespace(word_style_name="Normal")}

    document = render_nodes_to_document(nodes, styles, base_dir=tmp_path)

    assert len(document.inline_shapes) == 1


def test_render_nodes_to_document_uses_placeholder_for_missing_images(tmp_path: Path) -> None:
    nodes = [ImageBlockNode(raw_path="missing.png", alt_text="diagram")]
    styles = {"paragraph": SimpleNamespace(word_style_name="Normal")}

    document = render_nodes_to_document(nodes, styles, base_dir=tmp_path)

    assert document.paragraphs[0].text == "[Image not found: missing.png]"


def test_render_nodes_to_document_renders_blockquotes_as_prefixed_paragraphs() -> None:
    nodes = [QuoteNode(inlines=[TextSpan(text="Quoted line")])]
    styles = {"paragraph": SimpleNamespace(word_style_name="Normal")}

    document = render_nodes_to_document(nodes, styles)

    assert document.paragraphs[0].text == "> Quoted line"


def test_render_nodes_to_document_preserves_inline_runs_inside_blockquotes() -> None:
    nodes = [
        QuoteNode(
            inlines=[
                TextSpan(text="Quote "),
                TextSpan(text="bold", bold=True),
                TextSpan(text=" "),
                TextSpan(text="italic", italic=True),
                TextSpan(text=" "),
                TextSpan(text="code", code=True),
                TextSpan(text=" "),
                TextSpan(text="OpenAI", href="https://openai.com"),
            ]
        )
    ]
    styles = {"paragraph": SimpleNamespace(word_style_name="Normal")}

    document = render_nodes_to_document(nodes, styles)

    paragraph = document.paragraphs[0]
    assert paragraph.text == "> Quote bold italic code OpenAI"
    assert any(run.bold for run in paragraph.runs if run.text == "bold")
    assert any(run.italic for run in paragraph.runs if run.text == "italic")
    assert any(run.font.name == "Consolas" for run in paragraph.runs if run.text == "code")
    assert "hyperlink" in paragraph._p.xml


def test_render_nodes_to_document_renders_horizontal_rules_as_separator_paragraphs() -> None:
    nodes = [HorizontalRuleNode()]
    styles = {"paragraph": SimpleNamespace(word_style_name="Normal")}

    document = render_nodes_to_document(nodes, styles)

    assert document.paragraphs[0].text == "---"


def test_resolve_image_path_returns_absolute_path() -> None:
    image_path = resolve_image_path("tests/fixtures/example.png", Path("."))

    assert image_path.name == "example.png"
    assert image_path.is_absolute()


def test_resolve_image_path_rejects_blank_input() -> None:
    with pytest.raises(ValueError, match="blank"):
        resolve_image_path("   ", Path("."))


def test_resolve_image_path_rejects_absolute_path() -> None:
    with pytest.raises(ValueError, match="outside base_dir"):
        resolve_image_path(Path("C:/asset.png").as_posix(), Path("tests"))


def test_resolve_image_path_rejects_traversal_outside_base_dir() -> None:
    with pytest.raises(ValueError, match="outside base_dir"):
        resolve_image_path("../escape.png", Path("tests"))


def test_render_missing_image_in_lenient_mode_creates_text_placeholder() -> None:
    assert render_missing_image_placeholder("missing.png") == "[Image not found: missing.png]"


def test_run_render_returns_success_for_valid_inputs(tmp_path: Path) -> None:
    output = tmp_path / "out.docx"
    request = RenderRequest(
        input_path=Path("tests/fixtures/sample.md"),
        output_path=output,
        style_path=Path("tests/fixtures/style.yaml"),
        mode="strict",
        output_mode="human",
    )

    result = run_render(request)
    document = WordDocument(output)

    assert result.success is True
    assert result.output_path == output
    assert output.exists() is True
    assert [paragraph.text for paragraph in document.paragraphs] == [
        "Sample Title",
        "Plain paragraph with emphasis, link, and code.",
        "Line one\nLine two with diagram alt.",
        "☑ done",
        "☐ pending",
    ]
    second_runs = document.paragraphs[1].runs
    assert any(run.italic for run in second_runs if run.text == "emphasis")
    assert any(run.font.name == "Consolas" for run in second_runs if run.text == "code")
    assert "hyperlink" in document.paragraphs[1]._p.xml


def test_run_render_preserves_inline_strikethrough_to_docx(tmp_path: Path) -> None:
    markdown_path = tmp_path / "strike.md"
    markdown_path.write_text("Keep ~~remove~~ text\n", encoding="utf-8")
    output = tmp_path / "strike.docx"
    request = RenderRequest(
        input_path=markdown_path,
        output_path=output,
        style_path=Path("tests/fixtures/style.yaml"),
        mode="strict",
        output_mode="human",
    )

    result = run_render(request)
    document = WordDocument(output)

    assert result.success is True
    runs = document.paragraphs[0].runs
    assert any(run.font.strike for run in runs if run.text == "remove")


def test_run_render_supports_non_heading1_documents(tmp_path: Path) -> None:
    markdown_path = tmp_path / "subheading.md"
    markdown_path.write_text("## Subheading\n\nBody text.\n", encoding="utf-8")
    output = tmp_path / "subheading.docx"
    request = RenderRequest(
        input_path=markdown_path,
        output_path=output,
        style_path=Path("tests/fixtures/style.yaml"),
        mode="strict",
        output_mode="human",
    )

    result = run_render(request)
    document = WordDocument(output)

    assert result.success is True
    assert [paragraph.text for paragraph in document.paragraphs] == [
        "Subheading",
        "Body text.",
    ]
    assert document.paragraphs[0].style.name == "Heading 2"


def test_run_render_uses_heading3_style_name_when_falling_back_to_heading1_settings(tmp_path: Path) -> None:
    markdown_path = tmp_path / "third-level.md"
    markdown_path.write_text("### Third level\n", encoding="utf-8")
    output = tmp_path / "third-level.docx"
    request = RenderRequest(
        input_path=markdown_path,
        output_path=output,
        style_path=Path("tests/fixtures/style.yaml"),
        mode="strict",
        output_mode="human",
    )

    result = run_render(request)
    document = WordDocument(output)

    assert result.success is True
    assert document.paragraphs[0].text == "Third level"
    assert document.paragraphs[0].style.name == "Heading 3"


def test_run_render_preserves_inline_heading_formatting_to_docx(tmp_path: Path) -> None:
    markdown_path = tmp_path / "heading-inline.md"
    markdown_path.write_text(
        "# Title **bold** *italic* ~~gone~~ `code` [OpenAI](https://openai.com)\n",
        encoding="utf-8",
    )
    output = tmp_path / "heading-inline.docx"
    request = RenderRequest(
        input_path=markdown_path,
        output_path=output,
        style_path=Path("tests/fixtures/style.yaml"),
        mode="strict",
        output_mode="human",
    )

    result = run_render(request)
    document = WordDocument(output)

    assert result.success is True
    paragraph = document.paragraphs[0]
    assert paragraph.text == "Title bold italic gone code OpenAI"
    assert paragraph.style.name == "Heading 1"
    assert any(run.bold for run in paragraph.runs if run.text == "bold")
    assert any(run.italic for run in paragraph.runs if run.text == "italic")
    assert any(run.font.strike for run in paragraph.runs if run.text == "gone")
    assert any(run.font.name == "Consolas" for run in paragraph.runs if run.text == "code")
    assert "hyperlink" in paragraph._p.xml


def test_run_render_allows_heading_free_documents_without_heading1_style(tmp_path: Path) -> None:
    markdown_path = tmp_path / "no-headings.md"
    markdown_path.write_text("Paragraph only.\n", encoding="utf-8")
    style_path = tmp_path / "style.yaml"
    style_path.write_text(
        """\
meta:
  name: no-headings
  version: 1

defaults:
  font_name: Times New Roman
  font_size: 12

elements:
  paragraph:
    based_on: Normal
    font_size: 12
""",
        encoding="utf-8",
    )
    output = tmp_path / "no-headings.docx"
    request = RenderRequest(
        input_path=markdown_path,
        output_path=output,
        style_path=style_path,
        mode="strict",
        output_mode="human",
    )

    result = run_render(request)
    document = WordDocument(output)

    assert result.success is True
    assert [paragraph.text for paragraph in document.paragraphs] == ["Paragraph only."]


def test_run_render_renders_markdown_tables_to_docx(tmp_path: Path) -> None:
    markdown_path = tmp_path / "table.md"
    markdown_path.write_text(
        "\n".join(
            [
                "| Name | Value |",
                "| --- | --- |",
                "| Alpha | 1 |",
                "| Beta | 2 |",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    output = tmp_path / "table.docx"
    request = RenderRequest(
        input_path=markdown_path,
        output_path=output,
        style_path=Path("tests/fixtures/style.yaml"),
        mode="strict",
        output_mode="human",
    )

    result = run_render(request)
    document = WordDocument(output)

    assert result.success is True
    assert len(document.tables) == 1
    table = document.tables[0]
    assert table.cell(0, 0).text == "Name"
    assert table.cell(1, 0).text == "Alpha"
    assert table.cell(2, 1).text == "2"


def test_run_render_preserves_inline_table_cell_formatting_to_docx(tmp_path: Path) -> None:
    markdown_path = tmp_path / "table-inline.md"
    markdown_path.write_text(
        "\n".join(
            [
                "| Name | Value |",
                "| --- | --- |",
                "| **Alpha** | `42` [ref](https://example.com) |",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    output = tmp_path / "table-inline.docx"
    request = RenderRequest(
        input_path=markdown_path,
        output_path=output,
        style_path=Path("tests/fixtures/style.yaml"),
        mode="strict",
        output_mode="human",
    )

    result = run_render(request)
    document = WordDocument(output)

    assert result.success is True
    table = document.tables[0]
    first_cell_runs = table.cell(1, 0).paragraphs[0].runs
    second_cell_paragraph = table.cell(1, 1).paragraphs[0]
    assert table.cell(1, 0).text == "Alpha"
    assert table.cell(1, 1).text == "42 ref"
    assert any(run.bold for run in first_cell_runs if run.text == "Alpha")
    assert any(run.font.name == "Consolas" for run in second_cell_paragraph.runs if run.text == "42")
    assert "hyperlink" in second_cell_paragraph._p.xml


def test_run_render_renders_markdown_code_blocks_to_docx(tmp_path: Path) -> None:
    markdown_path = tmp_path / "code.md"
    markdown_path.write_text(
        "\n".join(
            [
                "```python",
                "print('hello')",
                "print('world')",
                "```",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    output = tmp_path / "code.docx"
    request = RenderRequest(
        input_path=markdown_path,
        output_path=output,
        style_path=Path("tests/fixtures/style.yaml"),
        mode="strict",
        output_mode="human",
    )

    result = run_render(request)
    document = WordDocument(output)

    assert result.success is True
    assert len(document.tables) == 1
    assert document.tables[0].cell(0, 0).text == "print('hello')\nprint('world')\n"


def test_run_render_renders_markdown_images_to_docx(tmp_path: Path) -> None:
    image_path = tmp_path / "diagram.png"
    image_path.write_bytes(_PNG_1X1)
    markdown_path = tmp_path / "image.md"
    markdown_path.write_text(f"![diagram]({image_path.name})\n", encoding="utf-8")
    output = tmp_path / "image.docx"
    request = RenderRequest(
        input_path=markdown_path,
        output_path=output,
        style_path=Path("tests/fixtures/style.yaml"),
        mode="strict",
        output_mode="human",
    )

    result = run_render(request)
    document = WordDocument(output)

    assert result.success is True
    assert len(document.inline_shapes) == 1


def test_run_render_renders_blockquotes_to_docx(tmp_path: Path) -> None:
    markdown_path = tmp_path / "quote.md"
    markdown_path.write_text("> Quoted line\n", encoding="utf-8")
    output = tmp_path / "quote.docx"
    request = RenderRequest(
        input_path=markdown_path,
        output_path=output,
        style_path=Path("tests/fixtures/style.yaml"),
        mode="strict",
        output_mode="human",
    )

    result = run_render(request)
    document = WordDocument(output)

    assert result.success is True
    assert [paragraph.text for paragraph in document.paragraphs] == ["> Quoted line"]


def test_run_render_preserves_inline_blockquote_formatting_to_docx(tmp_path: Path) -> None:
    markdown_path = tmp_path / "quote-inline.md"
    markdown_path.write_text(
        "> Quote **bold** *italic* `code` [OpenAI](https://openai.com)\n",
        encoding="utf-8",
    )
    output = tmp_path / "quote-inline.docx"
    request = RenderRequest(
        input_path=markdown_path,
        output_path=output,
        style_path=Path("tests/fixtures/style.yaml"),
        mode="strict",
        output_mode="human",
    )

    result = run_render(request)
    document = WordDocument(output)

    assert result.success is True
    paragraph = document.paragraphs[0]
    assert paragraph.text == "> Quote bold italic code OpenAI"
    assert any(run.bold for run in paragraph.runs if run.text == "bold")
    assert any(run.italic for run in paragraph.runs if run.text == "italic")
    assert any(run.font.name == "Consolas" for run in paragraph.runs if run.text == "code")
    assert "hyperlink" in paragraph._p.xml


def test_run_render_renders_horizontal_rules_to_docx(tmp_path: Path) -> None:
    markdown_path = tmp_path / "rule.md"
    markdown_path.write_text("---\n", encoding="utf-8")
    output = tmp_path / "rule.docx"
    request = RenderRequest(
        input_path=markdown_path,
        output_path=output,
        style_path=Path("tests/fixtures/style.yaml"),
        mode="strict",
        output_mode="human",
    )

    result = run_render(request)
    document = WordDocument(output)

    assert result.success is True
    assert [paragraph.text for paragraph in document.paragraphs] == ["---"]


def test_run_render_preserves_nested_markdown_lists_to_docx(tmp_path: Path) -> None:
    markdown_path = tmp_path / "nested-list.md"
    markdown_path.write_text(
        "\n".join(
            [
                "- Parent",
                "  - Child",
                "",
                "1. Ordered parent",
                "   1. Ordered child",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    output = tmp_path / "nested-list.docx"
    request = RenderRequest(
        input_path=markdown_path,
        output_path=output,
        style_path=Path("tests/fixtures/style.yaml"),
        mode="strict",
        output_mode="human",
    )

    result = run_render(request)
    document = WordDocument(output)

    assert result.success is True
    assert [paragraph.text for paragraph in document.paragraphs] == [
        "- Parent",
        "  - Child",
        "1. Ordered parent",
        "  1. Ordered child",
    ]


def test_run_render_preserves_inline_list_item_formatting_to_docx(tmp_path: Path) -> None:
    markdown_path = tmp_path / "list-inline.md"
    markdown_path.write_text(
        "- Plain **bold** *italic* `code` [OpenAI](https://openai.com)\n",
        encoding="utf-8",
    )
    output = tmp_path / "list-inline.docx"
    request = RenderRequest(
        input_path=markdown_path,
        output_path=output,
        style_path=Path("tests/fixtures/style.yaml"),
        mode="strict",
        output_mode="human",
    )

    result = run_render(request)
    document = WordDocument(output)

    assert result.success is True
    paragraph = document.paragraphs[0]
    assert paragraph.text == "- Plain bold italic code OpenAI"
    assert any(run.bold for run in paragraph.runs if run.text == "bold")
    assert any(run.italic for run in paragraph.runs if run.text == "italic")
    assert any(run.font.name == "Consolas" for run in paragraph.runs if run.text == "code")
    assert "hyperlink" in paragraph._p.xml


def test_run_render_preserves_task_list_inline_formatting_to_docx(tmp_path: Path) -> None:
    markdown_path = tmp_path / "task-list-inline.md"
    markdown_path.write_text(
        "- [x] **done** ~~old~~ `code` [ref](https://example.com)\n",
        encoding="utf-8",
    )
    output = tmp_path / "task-list-inline.docx"
    request = RenderRequest(
        input_path=markdown_path,
        output_path=output,
        style_path=Path("tests/fixtures/style.yaml"),
        mode="strict",
        output_mode="human",
    )

    result = run_render(request)
    document = WordDocument(output)

    assert result.success is True
    paragraph = document.paragraphs[0]
    assert paragraph.text == "☑ done old code ref"
    assert any(run.bold for run in paragraph.runs if run.text == "done")
    assert any(run.font.strike for run in paragraph.runs if run.text == "old")
    assert any(run.font.name == "Consolas" for run in paragraph.runs if run.text == "code")
    assert "hyperlink" in paragraph._p.xml


def test_render_nodes_to_document_rejects_unsupported_nodes() -> None:
    styles = {"paragraph": SimpleNamespace(word_style_name="Normal")}

    with pytest.raises(TypeError, match="Unsupported node type"):
        render_nodes_to_document([SimpleNamespace(type="mystery")], styles)


def test_render_nodes_to_document_preserves_mixed_block_order() -> None:
    nodes = [
        HeadingNode(level=1, inlines=[TextSpan(text="Title")]),
        ParagraphNode(inlines=[TextSpan(text="Paragraph A")]),
        TableNode(
            headers=[[TextSpan(text="Header")]],
            rows=[[[TextSpan(text="Value")]]],
        ),
        ParagraphNode(inlines=[TextSpan(text="Paragraph B")]),
        CodeBlockNode(content="code", language="python"),
        ImageBlockNode(raw_path="missing.png"),
        ParagraphNode(inlines=[TextSpan(text="Paragraph C")]),
    ]
    styles = {
        "heading1": SimpleNamespace(word_style_name="Heading 1", resolved_properties={}),
        "paragraph": SimpleNamespace(word_style_name="Normal", resolved_properties={}),
        "code_block": SimpleNamespace(word_style_name="Normal", resolved_properties={"font_name": "Consolas"}),
        "image": SimpleNamespace(word_style_name="Normal", resolved_properties={}),
    }

    document = render_nodes_to_document(nodes, styles)
    body = document.element.body
    texts = []
    for child in body:
        if child.tag == qn("w:p"):
            texts.append("".join(t.text for t in child.iter(qn("w:t"))))
        elif child.tag == qn("w:tbl"):
            first_cell = child.find(".//w:t", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
            texts.append(first_cell.text if first_cell is not None else "")

    assert texts == [
        "Title",
        "Paragraph A",
        "Header",
        "Paragraph B",
        "code",
        "[Image not found: missing.png]",
        "Paragraph C",
    ]


def test_render_nodes_to_document_allows_builtin_template_style_without_template() -> None:
    nodes = [
        HeadingNode(level=1, inlines=[TextSpan(text="Title")]),
        ParagraphNode(inlines=[TextSpan(text="Body")]),
    ]
    styles = {
        "heading1": SimpleNamespace(
            word_style_name="Heading 1",
            resolved_properties={"font_size": 18},
        ),
        "paragraph": SimpleNamespace(
            word_style_name="Normal",
            resolved_properties={"font_size": 12},
        ),
    }

    document = render_nodes_to_document(nodes, styles)

    assert document.paragraphs[0].text == "Title"
    assert document.paragraphs[0].style.name == "Heading 1"
    assert document.paragraphs[1].text == "Body"
    assert document.paragraphs[1].style.name == "Normal"


def test_render_nodes_to_document_applies_line_spacing_and_no_first_line_indent() -> None:
    nodes = [
        ParagraphNode(inlines=[TextSpan(text="Body text")]),
    ]
    styles = {
        "paragraph": SimpleNamespace(
            word_style_name="Normal",
            resolved_properties={
                "font_ascii": "Times New Roman",
                "font_east_asia": "宋体",
                "font_size": 12,
                "line_spacing": 1.0,
                "first_line_indent_pt": 0,
            },
        ),
    }

    document = render_nodes_to_document(nodes, styles)

    paragraph = document.paragraphs[0]
    run = paragraph.runs[0]
    assert paragraph.paragraph_format.line_spacing == 1.0
    assert paragraph.paragraph_format.first_line_indent is not None
    assert paragraph.paragraph_format.first_line_indent.pt == 0
    assert run.font.name == "Times New Roman"
    assert 'w:eastAsia="宋体"' in run._element.xml
