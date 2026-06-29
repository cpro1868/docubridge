from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from docubridge.core.template_bridge import TemplateView, load_template_view


def test_template_view_get_style_returns_named_style() -> None:
    template = TemplateView(
        available_styles={
            "Heading 1": {"font_name": "Arial", "font_size": 16},
        }
    )

    assert template.get_style("Heading 1") == {"font_name": "Arial", "font_size": 16}


def test_template_view_get_style_returns_copy() -> None:
    style = {"font_name": "Arial", "font_size": 16}
    template = TemplateView(available_styles={"Heading 1": style})

    resolved = template.get_style("Heading 1")
    resolved["font_size"] = 18

    assert template.available_styles["Heading 1"]["font_size"] == 16
    assert style["font_size"] == 16


def test_load_template_view_extracts_named_style_properties(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    document = Document()
    document.styles["Heading 1"].font.name = "Courier New"
    document.styles["Normal"].font.name = "Arial"
    document.save(template_path)

    template = load_template_view(template_path)

    assert template.get_style("Heading 1")["font_name"] == "Courier New"
    assert template.get_style("Normal")["font_name"] == "Arial"


def test_load_template_view_extracts_font_slots_and_paragraph_properties(tmp_path: Path) -> None:
    template_path = tmp_path / "template-fonts.docx"
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.first_line_indent = Pt(21)
    normal.paragraph_format.space_after = Pt(6)
    document.save(template_path)

    template = load_template_view(template_path)

    normal_style = template.get_style("Normal")
    assert normal_style["font_ascii"] == "Times New Roman"
    assert normal_style["font_east_asia"] == "宋体"
    assert normal_style["first_line_indent_pt"] == 21
    assert normal_style["space_after_pt"] == 6


def test_load_template_view_extracts_numbering_definitions(tmp_path: Path) -> None:
    template_path = tmp_path / "template-numbering.docx"
    document = Document()
    document.add_paragraph("One", style="List Number")
    document.add_paragraph("Two", style="List Number 2")
    document.save(template_path)

    template = load_template_view(template_path)

    assert template.available_numberings
    assert "List Number" in template.style_numbering_map
