from __future__ import annotations

import base64
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches

from docubridge.core.docx_ingest import parse_docx_file


_PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+aRX0AAAAASUVORK5CYII="
)


def _write_text_list_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Sample Title", style="Heading 1")
    document.add_paragraph("Plain paragraph.")
    document.add_paragraph("First bullet", style="List Bullet")
    document.add_paragraph("Second bullet", style="List Bullet")
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _write_rich_text_docx(path: Path) -> None:
    document = Document()
    heading = document.add_paragraph(style="Heading 1")
    heading.add_run("Title ")
    heading.add_run("Bold").bold = True
    heading.add_run(" ")
    heading.add_run("Italic").italic = True
    heading.add_run(" ")
    heading.add_run("Gone").font.strike = True
    heading.add_run(" ")
    _append_hyperlink(heading, "ref", "https://example.com")

    paragraph = document.add_paragraph()
    paragraph.add_run("Body ")
    paragraph.add_run("Bold").bold = True
    paragraph.add_run(" ")
    paragraph.add_run("Italic").italic = True
    paragraph.add_run(" ")
    paragraph.add_run("Gone").font.strike = True
    paragraph.add_run(" ")
    _append_hyperlink(paragraph, "ref", "https://example.com")

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _write_code_text_docx(path: Path) -> None:
    document = Document()
    heading = document.add_paragraph(style="Heading 1")
    heading.add_run("Title ")
    code_run = heading.add_run("code")
    code_run.font.name = "Consolas"

    paragraph = document.add_paragraph()
    paragraph.add_run("Body ")
    body_code = paragraph.add_run("code")
    body_code.font.name = "Consolas"

    table = document.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "Value"
    cell = table.cell(1, 0)
    cell.text = ""
    cell_paragraph = cell.paragraphs[0]
    cell_paragraph.add_run("Cell ")
    cell_code = cell_paragraph.add_run("code")
    cell_code.font.name = "Consolas"

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _write_fragmented_bold_docx(path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    first = paragraph.add_run("1")
    first.bold = True
    second = paragraph.add_run("、")
    second.bold = True
    third = paragraph.add_run("故事发生地点：")
    third.bold = True
    paragraph.add_run("卡拉尔星东半球西大陆")

    table = document.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "Value"
    cell = table.cell(1, 0)
    cell.text = ""
    cell_paragraph = cell.paragraphs[0]
    cell_first = cell_paragraph.add_run("10 ")
    cell_first.bold = True
    cell_second = cell_paragraph.add_run("天")
    cell_second.bold = True

    paragraph_with_punctuation = document.add_paragraph()
    punct_first = paragraph_with_punctuation.add_run("2")
    punct_first.bold = True
    punct_second = paragraph_with_punctuation.add_run("、")
    punct_second.bold = True
    punct_third = paragraph_with_punctuation.add_run("时代科技水平")
    punct_third.bold = True
    paragraph_with_punctuation.add_run("：整体不如现代地球")

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _write_table_docx(path: Path) -> None:
    document = Document()
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "1"
    table.cell(2, 0).text = "Beta"
    table.cell(2, 1).text = "2"
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _append_hyperlink(paragraph, text: str, href: str) -> None:
    part = paragraph.part
    rel_id = part.relate_to(href, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _get_or_add_paragraph_properties(paragraph):
    properties = paragraph._p.pPr
    if properties is None:
        properties = OxmlElement("w:pPr")
        paragraph._p.insert(0, properties)
    return properties


def _set_paragraph_numbering(paragraph, *, num_id: int, ilvl: int = 0) -> None:
    properties = _get_or_add_paragraph_properties(paragraph)
    num_pr = properties.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        properties.append(num_pr)

    ilvl_element = num_pr.find(qn("w:ilvl"))
    if ilvl_element is None:
        ilvl_element = OxmlElement("w:ilvl")
        num_pr.append(ilvl_element)
    ilvl_element.set(qn("w:val"), str(ilvl))

    num_id_element = num_pr.find(qn("w:numId"))
    if num_id_element is None:
        num_id_element = OxmlElement("w:numId")
        num_pr.append(num_id_element)
    num_id_element.set(qn("w:val"), str(num_id))


def _style_numbering_reference(style) -> tuple[int, int] | None:
    current = style
    while current is not None:
        properties = getattr(current.element, "pPr", None)
        num_pr = getattr(properties, "numPr", None)
        if num_pr is not None and getattr(num_pr, "numId", None) is not None:
            ilvl = getattr(num_pr, "ilvl", None)
            return int(num_pr.numId.val), int(ilvl.val if ilvl is not None else 0)
        current = current.base_style
    return None


def _paragraph_numbering_reference(paragraph) -> tuple[int, int]:
    properties = paragraph._p.pPr
    num_pr = getattr(properties, "numPr", None) if properties is not None else None
    if num_pr is not None and getattr(num_pr, "numId", None) is not None:
        ilvl = getattr(num_pr, "ilvl", None)
        return int(num_pr.numId.val), int(ilvl.val if ilvl is not None else 0)

    reference = _style_numbering_reference(paragraph.style)
    if reference is None:
        raise AssertionError("Paragraph has no numbering reference")
    return reference


def _set_numbering_start_override(document: Document, *, num_id: int, ilvl: int, start_at: int) -> None:
    numbering = document.part.numbering_part.element
    num_nodes = numbering.xpath(f"./w:num[@w:numId='{num_id}']")
    if not num_nodes:
        raise AssertionError(f"Missing numbering definition for numId={num_id}")
    num_node = num_nodes[0]

    override_nodes = num_node.xpath(f"./w:lvlOverride[@w:ilvl='{ilvl}']")
    if override_nodes:
        override_node = override_nodes[0]
    else:
        override_node = OxmlElement("w:lvlOverride")
        override_node.set(qn("w:ilvl"), str(ilvl))
        num_node.append(override_node)

    start_override = override_node.find(qn("w:startOverride"))
    if start_override is None:
        start_override = OxmlElement("w:startOverride")
        override_node.append(start_override)
    start_override.set(qn("w:val"), str(start_at))


def _write_rich_table_docx(path: Path) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"

    first_cell = table.cell(1, 0)
    first_cell.text = ""
    first_paragraph = first_cell.paragraphs[0]
    first_paragraph.add_run("Alpha").bold = True
    first_paragraph.add_run(" ")
    first_paragraph.add_run("Italic").italic = True

    second_cell = table.cell(1, 1)
    second_cell.text = ""
    second_paragraph = second_cell.paragraphs[0]
    second_paragraph.add_run("See ")
    _append_hyperlink(second_paragraph, "ref", "https://example.com")

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _write_table_with_multiline_cell_docx(path: Path) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "Line one\nLine two"
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _write_ordered_docx(path: Path, image_path: Path) -> None:
    image_path.parent.mkdir(parents=True, exist_ok=True)
    image_path.write_bytes(_PNG_1X1)
    document = Document()
    document.add_paragraph("Before image")
    document.add_picture(str(image_path))
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Key"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Order"
    table.cell(1, 1).text = "Preserved"
    document.add_paragraph("After table")
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _write_nested_list_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Top", style="List Bullet")
    nested = document.add_paragraph("Nested", style="List Bullet")
    nested.paragraph_format.left_indent = Inches(0.5)
    document.add_paragraph("Top 2", style="List Bullet")
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _write_numbered_list_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("First", style="List Number")
    document.add_paragraph("Second", style="List Number")
    nested = document.add_paragraph("Nested", style="List Number")
    nested.paragraph_format.left_indent = Inches(0.5)
    document.add_paragraph("Third", style="List Number")
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _write_numbered_start_override_docx(path: Path) -> None:
    document = Document()
    first = document.add_paragraph("Five", style="List Number")
    second = document.add_paragraph("Six", style="List Number")
    num_id, ilvl = _paragraph_numbering_reference(first)
    _set_numbering_start_override(document, num_id=num_id, ilvl=ilvl, start_at=5)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _write_numbered_heading_docx(path: Path) -> None:
    document = Document()
    reference = document.add_paragraph("Reference", style="List Number")
    num_id, ilvl = _paragraph_numbering_reference(reference)
    heading = document.add_paragraph("Heading With Number", style="Heading 1")
    _set_paragraph_numbering(heading, num_id=num_id, ilvl=ilvl)
    reference._element.getparent().remove(reference._element)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _write_numbered_headings_with_body_docx(path: Path) -> None:
    document = Document()
    reference = document.add_paragraph("Reference", style="List Number")
    num_id, ilvl = _paragraph_numbering_reference(reference)

    first_heading = document.add_paragraph("First Section", style="Heading 1")
    _set_paragraph_numbering(first_heading, num_id=num_id, ilvl=ilvl)
    document.add_paragraph("Body under first section.")

    second_heading = document.add_paragraph("Second Section", style="Heading 1")
    _set_paragraph_numbering(second_heading, num_id=num_id, ilvl=ilvl)
    document.add_paragraph("Body under second section.")

    reference._element.getparent().remove(reference._element)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _write_numbered_headings_with_table_docx(path: Path) -> None:
    document = Document()
    reference = document.add_paragraph("Reference", style="List Number")
    num_id, ilvl = _paragraph_numbering_reference(reference)

    first_heading = document.add_paragraph("First Section", style="Heading 1")
    _set_paragraph_numbering(first_heading, num_id=num_id, ilvl=ilvl)

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Key"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "1"

    second_heading = document.add_paragraph("Second Section", style="Heading 1")
    _set_paragraph_numbering(second_heading, num_id=num_id, ilvl=ilvl)

    reference._element.getparent().remove(reference._element)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _write_release_case_docx(path: Path, image_path: Path) -> None:
    image_path.parent.mkdir(parents=True, exist_ok=True)
    image_path.write_bytes(_PNG_1X1)

    document = Document()
    heading = document.add_paragraph(style="Heading 1")
    heading.add_run("Release ")
    heading.add_run("Bold").bold = True

    paragraph = document.add_paragraph()
    paragraph.add_run("Body ")
    paragraph.add_run("Italic").italic = True
    paragraph.add_run(" ")
    code_run = paragraph.add_run("code")
    code_run.font.name = "Consolas"

    document.add_paragraph("Top bullet", style="List Bullet")
    nested = document.add_paragraph("Nested bullet", style="List Bullet")
    nested.paragraph_format.left_indent = Inches(0.5)

    document.add_picture(str(image_path))

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "1"

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def test_parse_docx_file_extracts_headings_paragraphs_and_lists(tmp_path: Path) -> None:
    input_path = tmp_path / "sample.docx"
    _write_text_list_docx(input_path)

    content = parse_docx_file(input_path)

    assert "# Sample Title" in content
    assert "Plain paragraph." in content
    assert "- First bullet" in content
    assert "- Second bullet" in content


def test_parse_docx_file_preserves_rich_text_in_headings_and_paragraphs(tmp_path: Path) -> None:
    input_path = tmp_path / "rich-text.docx"
    _write_rich_text_docx(input_path)

    content = parse_docx_file(input_path)

    assert "# Title **Bold** *Italic* ~~Gone~~ [ref](https://example.com)" in content
    assert "Body **Bold** *Italic* ~~Gone~~ [ref](https://example.com)" in content


def test_parse_docx_file_preserves_inline_code_in_headings_paragraphs_and_tables(tmp_path: Path) -> None:
    input_path = tmp_path / "code-text.docx"
    _write_code_text_docx(input_path)

    content = parse_docx_file(input_path)

    assert "# Title `code`" in content
    assert "Body `code`" in content
    assert "| Cell `code` |" in content


def test_parse_docx_file_merges_adjacent_bold_runs_into_single_markdown_span(
    tmp_path: Path,
) -> None:
    input_path = tmp_path / "fragmented-bold.docx"
    _write_fragmented_bold_docx(input_path)

    content = parse_docx_file(input_path)

    assert "**1、故事发生地点：** 卡拉尔星东半球西大陆" in content
    assert "**2、时代科技水平：** 整体不如现代地球" in content
    assert "**1****、****故事发生地点：**" not in content
    assert "| **10 天** |" in content
    assert "**10 ****天**" not in content


def test_parse_docx_file_extracts_tables_as_markdown(tmp_path: Path) -> None:
    input_path = tmp_path / "table.docx"
    _write_table_docx(input_path)

    content = parse_docx_file(input_path)

    assert "| Name | Value |" in content
    assert "| --- | --- |" in content
    assert "| Alpha | 1 |" in content
    assert "| Beta | 2 |" in content


def test_parse_docx_file_preserves_rich_text_inside_table_cells(tmp_path: Path) -> None:
    input_path = tmp_path / "rich-table.docx"
    _write_rich_table_docx(input_path)

    content = parse_docx_file(input_path)

    assert "| Name | Value |" in content
    assert "| **Alpha** *Italic* | See [ref](https://example.com) |" in content


def test_parse_docx_file_preserves_multiline_table_cells_as_single_markdown_row(
    tmp_path: Path,
) -> None:
    input_path = tmp_path / "multiline-table.docx"
    _write_table_with_multiline_cell_docx(input_path)

    content = parse_docx_file(input_path)

    assert "| Alpha | Line one<br>Line two |" in content
    assert "Line one\nLine two |" not in content


def test_parse_docx_file_preserves_image_and_table_order(tmp_path: Path) -> None:
    input_path = tmp_path / "ordered.docx"
    image_path = tmp_path / "ordered.png"
    media_dir = tmp_path / "media"
    _write_ordered_docx(input_path, image_path)

    content = parse_docx_file(input_path, media_dir=media_dir)

    before_index = content.index("Before image")
    image_index = content.index("![ordered](")
    table_index = content.index("| Key | Value |")
    after_index = content.index("After table")
    assert before_index < image_index < table_index < after_index
    assert list(media_dir.iterdir())


def test_parse_docx_file_preserves_simple_nested_lists(tmp_path: Path) -> None:
    input_path = tmp_path / "nested.docx"
    _write_nested_list_docx(input_path)

    content = parse_docx_file(input_path)

    assert "- Top" in content
    assert "  - Nested" in content
    assert "- Top 2" in content


def test_parse_docx_file_preserves_simple_numbered_list_progression(tmp_path: Path) -> None:
    input_path = tmp_path / "numbered.docx"
    _write_numbered_list_docx(input_path)

    content = parse_docx_file(input_path)

    assert "1. First" in content
    assert "2. Second" in content
    assert "  1. Nested" in content
    assert "3. Third" in content


def test_parse_docx_file_preserves_original_numbering_values(tmp_path: Path) -> None:
    input_path = tmp_path / "numbered-start.docx"
    _write_numbered_start_override_docx(input_path)

    content = parse_docx_file(input_path)

    assert "5. Five" in content
    assert "6. Six" in content


def test_parse_docx_file_preserves_numbered_headings(tmp_path: Path) -> None:
    input_path = tmp_path / "heading-numbered.docx"
    _write_numbered_heading_docx(input_path)

    content = parse_docx_file(input_path)

    assert "# 1. Heading With Number" in content


def test_parse_docx_file_keeps_heading_number_progression_across_body_paragraphs(
    tmp_path: Path,
) -> None:
    input_path = tmp_path / "heading-numbered-body.docx"
    _write_numbered_headings_with_body_docx(input_path)

    content = parse_docx_file(input_path)

    assert "# 1. First Section" in content
    assert "# 2. Second Section" in content


def test_parse_docx_file_keeps_heading_number_progression_across_tables(
    tmp_path: Path,
) -> None:
    input_path = tmp_path / "heading-numbered-table.docx"
    _write_numbered_headings_with_table_docx(input_path)

    content = parse_docx_file(input_path)

    assert "# 1. First Section" in content
    assert "# 2. Second Section" in content


def test_parse_docx_file_release_case_preserves_core_v1_structures(tmp_path: Path) -> None:
    input_path = tmp_path / "release.docx"
    image_path = tmp_path / "release.png"
    media_dir = tmp_path / "assets"
    _write_release_case_docx(input_path, image_path)

    content = parse_docx_file(input_path, media_dir=media_dir)

    assert "# Release **Bold**" in content
    assert "Body *Italic* `code`" in content
    assert "- Top bullet" in content
    assert "  - Nested bullet" in content
    assert "![" in content
    assert "assets/" in content
    assert "| Name | Value |" in content
    assert "| Alpha | 1 |" in content
